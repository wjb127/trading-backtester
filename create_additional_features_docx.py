"""
추가 기능 구현안 문서를 DOCX로 변환
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_additional_features_docx():
    # 문서 생성
    doc = Document()

    # 제목
    title = doc.add_heading('Trading Backtest 추가 기능 구현안', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 프로젝트 현재 상태
    doc.add_heading('프로젝트 현재 상태', 1)

    doc.add_heading('완료된 기능', 2)
    completed = [
        '백엔드 API 기본 구조 (FastAPI + Supabase REST API)',
        '프론트엔드 기본 구조 (Next.js 14 + TypeScript + Tailwind CSS)',
        'Supabase 데이터베이스 스키마 설계',
        '전략 관리 API (CRUD)',
        '기본 페이지 구조 (대시보드, 전략, 백테스트)'
    ]
    for item in completed:
        doc.add_paragraph(f'✅ {item}', style='List Bullet')

    doc.add_heading('미완성 기능 (Placeholder)', 2)
    incomplete = [
        '백테스트 실행 엔진',
        '시장 데이터 수집 및 저장',
        '성과 분석 및 통계',
        '차트 및 시각화'
    ]
    for item in incomplete:
        doc.add_paragraph(f'🔲 {item}', style='List Bullet')

    # 1단계: 백테스트 실행 엔진
    doc.add_page_break()
    doc.add_heading('1단계: 백테스트 실행 엔진', 1)

    doc.add_heading('1.1 전략 실행 엔진', 2)
    doc.add_paragraph('목표: 사용자가 작성한 전략 코드를 안전하게 실행하고 매매 신호를 생성')

    doc.add_heading('주요 기능', 3)
    features = [
        '코드 샌드박싱: RestrictedPython 사용하여 안전한 코드 실행',
        '지표 라이브러리: TA-Lib, pandas-ta 통합',
        '포지션 관리: 롱/숏 포지션, 레버리지, 수수료 계산',
        '슬리피지 모델링: 실제 거래와 유사한 환경 구현'
    ]
    for feature in features:
        doc.add_paragraph(feature, style='List Bullet')

    doc.add_heading('API 엔드포인트', 3)
    api_endpoints_1 = [
        'POST /api/v1/backtests',
        '  - Request Body: { strategy_id, symbol, start_date, end_date, initial_capital }',
        '  - Response: { backtest_id, status }',
        '',
        'GET /api/v1/backtests/{backtest_id}/status',
        '  - Response: { status, progress, current_step }',
        '',
        'GET /api/v1/backtests/{backtest_id}/result',
        '  - Response: { trades, metrics, equity_curve }'
    ]
    for endpoint in api_endpoints_1:
        if endpoint:
            doc.add_paragraph(endpoint, style='List Bullet')

    # 2단계: 시장 데이터 수집
    doc.add_page_break()
    doc.add_heading('2단계: 시장 데이터 수집', 1)

    doc.add_heading('2.1 데이터 소스 통합', 2)
    doc.add_paragraph('목표: 주식 및 암호화폐 과거 데이터 수집 및 저장')

    doc.add_heading('지원 데이터 소스', 3)

    doc.add_paragraph('주식 데이터:', style='List Bullet')
    stock_sources = [
        'Yahoo Finance API (yfinance)',
        'Alpha Vantage API',
        '한국투자증권 API (KIS Developers)'
    ]
    for source in stock_sources:
        p = doc.add_paragraph(source, style='List Bullet 2')

    doc.add_paragraph('암호화폐 데이터:', style='List Bullet')
    crypto_sources = [
        'Binance API',
        'Upbit API',
        'CoinGecko API'
    ]
    for source in crypto_sources:
        p = doc.add_paragraph(source, style='List Bullet 2')

    # 3단계: 성과 분석 및 통계
    doc.add_page_break()
    doc.add_heading('3단계: 성과 분석 및 통계', 1)

    doc.add_heading('3.1 백테스트 성과 지표', 2)
    doc.add_paragraph('목표: 백테스트 결과를 다양한 지표로 분석')

    doc.add_heading('주요 지표', 3)

    doc.add_paragraph('수익성 지표:', style='List Bullet')
    profit_metrics = [
        '총 수익률 (Total Return)',
        '연간 수익률 (Annual Return)',
        '최대 낙폭 (Maximum Drawdown)',
        '승률 (Win Rate)',
        '손익비 (Profit Factor)'
    ]
    for metric in profit_metrics:
        doc.add_paragraph(metric, style='List Bullet 2')

    doc.add_paragraph('리스크 지표:', style='List Bullet')
    risk_metrics = [
        '샤프 비율 (Sharpe Ratio)',
        '소르티노 비율 (Sortino Ratio)',
        '변동성 (Volatility)'
    ]
    for metric in risk_metrics:
        doc.add_paragraph(metric, style='List Bullet 2')

    doc.add_paragraph('거래 지표:', style='List Bullet')
    trade_metrics = [
        '총 거래 횟수',
        '평균 보유 기간',
        '평균 수익/손실'
    ]
    for metric in trade_metrics:
        doc.add_paragraph(metric, style='List Bullet 2')

    # 4단계: 차트 및 시각화
    doc.add_page_break()
    doc.add_heading('4단계: 차트 및 시각화', 1)

    doc.add_heading('4.1 프론트엔드 차트 라이브러리', 2)
    doc.add_paragraph('목표: 백테스트 결과를 시각적으로 표현')

    doc.add_heading('사용 라이브러리', 3)
    chart_libs = [
        'Chart.js 또는 Recharts: 성과 차트',
        'Lightweight Charts (TradingView): 캔들스틱 차트',
        'D3.js: 고급 시각화 (선택)'
    ]
    for lib in chart_libs:
        doc.add_paragraph(lib, style='List Bullet')

    doc.add_heading('구현할 차트', 3)
    charts = [
        '수익 곡선 (Equity Curve): 시간에 따른 포트폴리오 가치 변화, 벤치마크 비교',
        '매매 신호 차트: 캔들스틱 + 매수/매도 마커, 지표 오버레이',
        '낙폭 차트 (Drawdown Chart): 시간에 따른 낙폭 변화',
        '월별 수익률 히트맵: 월별, 연도별 수익률 분포'
    ]
    for chart in charts:
        doc.add_paragraph(chart, style='List Bullet')

    # 5단계: 전략 최적화
    doc.add_page_break()
    doc.add_heading('5단계: 전략 최적화', 1)

    doc.add_heading('5.1 파라미터 최적화', 2)
    doc.add_paragraph('목표: 전략의 최적 파라미터 자동 탐색')

    doc.add_heading('최적화 방법', 3)
    optimization_methods = [
        '그리드 서치 (Grid Search): 모든 파라미터 조합 탐색, 계산 비용 높음',
        '무작위 서치 (Random Search): 무작위 파라미터 샘플링, 그리드 서치보다 효율적',
        '베이지안 최적화 (Bayesian Optimization): 과거 결과를 바탕으로 다음 파라미터 선택, 가장 효율적'
    ]
    for method in optimization_methods:
        doc.add_paragraph(method, style='List Bullet')

    # 6단계: 실시간 모니터링
    doc.add_page_break()
    doc.add_heading('6단계: 실시간 모니터링 (선택)', 1)

    doc.add_heading('6.1 실시간 데이터 스트리밍', 2)
    doc.add_paragraph('목표: 실시간 시장 데이터 수집 및 전략 시뮬레이션')

    doc.add_heading('구현 방법', 3)
    streaming_methods = [
        'WebSocket: 실시간 데이터 수신',
        'Supabase Realtime: 데이터베이스 변경 사항 실시간 구독'
    ]
    for method in streaming_methods:
        doc.add_paragraph(method, style='List Bullet')

    # 7단계: 알림 및 리포트
    doc.add_heading('7단계: 알림 및 리포트', 1)

    doc.add_heading('7.1 백테스트 완료 알림', 2)
    doc.add_paragraph('목표: 백테스트 완료 시 사용자에게 알림')

    doc.add_heading('알림 방법', 3)
    notification_methods = [
        '이메일 (SendGrid, AWS SES)',
        '웹 푸시 알림',
        'Slack/Discord 웹훅'
    ]
    for method in notification_methods:
        doc.add_paragraph(method, style='List Bullet')

    doc.add_heading('7.2 PDF 리포트 생성', 2)
    doc.add_paragraph('목표: 백테스트 결과를 PDF로 내보내기')

    doc.add_heading('구현 라이브러리', 3)
    pdf_libs = [
        'reportlab (Python)',
        'jsPDF (JavaScript)'
    ]
    for lib in pdf_libs:
        doc.add_paragraph(lib, style='List Bullet')

    # 우선순위
    doc.add_page_break()
    doc.add_heading('우선순위', 1)

    doc.add_heading('높은 우선순위 (즉시 구현)', 2)
    high_priority = [
        '백테스트 실행 엔진 (1단계)',
        '시장 데이터 수집 (2단계)',
        '성과 분석 (3단계)'
    ]
    for item in high_priority:
        doc.add_paragraph(f'✅ {item}', style='List Bullet')

    doc.add_heading('중간 우선순위 (다음 단계)', 2)
    mid_priority = [
        '차트 및 시각화 (4단계)',
        '전략 최적화 (5단계)'
    ]
    for item in mid_priority:
        doc.add_paragraph(f'✅ {item}', style='List Bullet')

    doc.add_heading('낮은 우선순위 (향후 검토)', 2)
    low_priority = [
        '실시간 모니터링 (6단계)',
        '알림 및 리포트 (7단계)'
    ]
    for item in low_priority:
        doc.add_paragraph(f'🔲 {item}', style='List Bullet')

    # 기술 스택
    doc.add_page_break()
    doc.add_heading('기술 스택 추가 요구사항', 1)

    doc.add_heading('Python 패키지', 2)
    python_packages = [
        '데이터 처리: pandas, numpy',
        '기술 지표: TA-Lib, pandas-ta',
        '데이터 수집: yfinance, python-binance, ccxt',
        '백테스팅: backtrader 또는 직접 구현',
        '최적화: scikit-optimize, optuna',
        '코드 샌드박싱: RestrictedPython',
        '비동기 작업: celery, redis'
    ]
    for package in python_packages:
        doc.add_paragraph(package, style='List Bullet')

    doc.add_heading('프론트엔드 패키지', 2)
    frontend_packages = [
        'chart.js, react-chartjs-2',
        'lightweight-charts',
        'date-fns',
        'recharts'
    ]
    for package in frontend_packages:
        doc.add_paragraph(package, style='List Bullet')

    # 개발 로드맵
    doc.add_page_break()
    doc.add_heading('개발 로드맵', 1)

    doc.add_heading('Phase 1: 코어 기능 (2-3주)', 2)
    phase1 = [
        '[x] 프로젝트 기본 구조',
        '[ ] 백테스트 엔진 구현',
        '[ ] 데이터 수집 API',
        '[ ] 기본 성과 분석'
    ]
    for item in phase1:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Phase 2: 시각화 (1-2주)', 2)
    phase2 = [
        '[ ] 차트 컴포넌트 개발',
        '[ ] 대시보드 개선',
        '[ ] 백테스트 결과 페이지'
    ]
    for item in phase2:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Phase 3: 고급 기능 (2-3주)', 2)
    phase3 = [
        '[ ] 파라미터 최적화',
        '[ ] 전략 비교 기능',
        '[ ] PDF 리포트 생성'
    ]
    for item in phase3:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Phase 4: 배포 및 최적화 (1주)', 2)
    phase4 = [
        '[ ] Fly.io 배포',
        '[ ] 성능 최적화',
        '[ ] 모니터링 설정'
    ]
    for item in phase4:
        doc.add_paragraph(item, style='List Bullet')

    # 주의사항
    doc.add_page_break()
    doc.add_heading('라이선스 및 주의사항', 1)

    doc.add_heading('주의사항', 2)
    warnings = [
        '백테스트 결과는 과거 데이터 기반이며 미래 수익을 보장하지 않음',
        '실제 거래 시 슬리피지, 수수료 등 추가 비용 고려 필요',
        '과최적화(overfitting) 주의'
    ]
    for warning in warnings:
        doc.add_paragraph(warning, style='List Bullet')

    doc.add_heading('라이선스', 2)
    license_info = [
        '개인 사용 목적',
        '상업적 사용 시 별도 라이선스 필요'
    ]
    for info in license_info:
        doc.add_paragraph(info, style='List Bullet')

    # 문서 정보
    doc.add_paragraph()
    doc.add_paragraph('문서 작성일: 2025-01-15')
    doc.add_paragraph('버전: 1.0')

    # 저장
    doc.save('추가기능_구현안.docx')
    print("✅ 추가기능_구현안.docx 파일이 생성되었습니다!")

if __name__ == "__main__":
    create_additional_features_docx()
