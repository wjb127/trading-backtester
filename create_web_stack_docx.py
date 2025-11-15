#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
웹 배포 기술 스택 DOCX 문서 생성 스크립트
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_web_stack_doc():
    doc = Document()

    # 제목
    title = doc.add_heading('주식/코인 트레이딩 백테스팅 프로그램', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading('웹 배포 기술 스택 문서', 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # 1. 프로젝트 개요
    doc.add_heading('1. 프로젝트 개요', 1)
    doc.add_paragraph(
        '백테스팅 프로그램을 웹 애플리케이션으로 구현하여 '
        'fly.io에 모노레포 구조로 배포합니다. '
        'PostgreSQL을 사용하며 향후 Supabase로 마이그레이션 예정입니다.'
    )

    # 2. 전체 기술 스택
    doc.add_heading('2. 전체 기술 스택', 1)

    doc.add_heading('2.1 백엔드', 2)
    backend_stack = [
        '프레임워크: FastAPI (Python 3.11+)',
        'ORM: SQLAlchemy 2.0 + Alembic (마이그레이션)',
        '인증: JWT (python-jose) + bcrypt',
        '비동기: asyncio, asyncpg',
        '작업 큐: Celery + Redis (백테스팅 장시간 작업용)',
        '데이터 처리: pandas, numpy, TA-Lib',
        'API 문서: Swagger/OpenAPI (FastAPI 자동 생성)'
    ]
    for item in backend_stack:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('2.2 프론트엔드', 2)
    frontend_stack = [
        '프레임워크: Next.js 14+ (App Router)',
        '언어: TypeScript',
        '상태관리: Zustand 또는 Tanstack Query',
        'UI 라이브러리: shadcn/ui + Tailwind CSS',
        '차트: Recharts 또는 Plotly.js',
        '데이터 테이블: AG Grid 또는 TanStack Table',
        'API 통신: Axios 또는 Fetch API',
        '폼 관리: React Hook Form + Zod (validation)'
    ]
    for item in frontend_stack:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('2.3 데이터베이스', 2)
    db_stack = [
        '현재: PostgreSQL 16',
        '향후: Supabase (PostgreSQL + Auth + Storage + Realtime)',
        '캐싱: Redis (세션, 캐시, Celery 브로커)'
    ]
    for item in db_stack:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('2.4 인프라 & 배포', 2)
    infra_stack = [
        '배포 플랫폼: fly.io',
        '컨테이너: Docker + Docker Compose',
        '프록시: Nginx (프론트/백엔드 라우팅)',
        'CI/CD: GitHub Actions',
        '모니터링: Sentry (에러 트래킹)'
    ]
    for item in infra_stack:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('2.5 개발 도구', 2)
    dev_tools = [
        '버전 관리: Git',
        'Backend 패키지 매니저: Poetry 또는 pip-tools',
        'Frontend 패키지 매니저: pnpm',
        'Backend 코드 품질: black, flake8, mypy',
        'Frontend 코드 품질: ESLint, Prettier',
        'Backend 테스팅: pytest, pytest-asyncio',
        'Frontend 테스팅: Vitest, React Testing Library'
    ]
    for item in dev_tools:
        doc.add_paragraph(item, style='List Bullet')

    # 3. 모노레포 구조
    doc.add_heading('3. 모노레포 구조', 1)

    structure = """trading-backtest/
├── backend/                      # FastAPI 백엔드
│   ├── app/
│   │   ├── main.py              # FastAPI 앱 진입점
│   │   ├── config.py            # 환경 설정
│   │   ├── database.py          # DB 연결
│   │   ├── api/v1/              # API 라우터
│   │   ├── models/              # SQLAlchemy 모델
│   │   ├── schemas/             # Pydantic 스키마
│   │   ├── services/            # 비즈니스 로직
│   │   ├── core/                # 백테스팅 핵심
│   │   ├── tasks/               # Celery 작업
│   │   └── utils/               # 유틸리티
│   ├── alembic/                 # DB 마이그레이션
│   ├── tests/
│   ├── Dockerfile
│   └── requirements.txt
│
├── frontend/                     # Next.js 프론트엔드
│   ├── src/
│   │   ├── app/                 # Next.js App Router
│   │   ├── components/          # React 컴포넌트
│   │   ├── lib/                 # 유틸리티
│   │   ├── hooks/               # Custom Hooks
│   │   ├── store/               # 상태 관리
│   │   └── types/               # TypeScript 타입
│   ├── Dockerfile
│   ├── next.config.js
│   └── package.json
│
├── nginx/                        # Nginx 설정
├── docker/                       # Docker 설정
├── scripts/                      # 유틸리티 스크립트
├── .github/workflows/            # CI/CD
├── fly.toml                      # Fly.io 설정
└── README.md"""

    p = doc.add_paragraph(structure)
    p.style = 'No Spacing'
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(8)

    # 4. API 설계
    doc.add_heading('4. API 설계', 1)

    api_sections = [
        ('4.1 인증 API', [
            'POST   /api/v1/auth/register      # 회원가입',
            'POST   /api/v1/auth/login         # 로그인',
            'POST   /api/v1/auth/refresh       # 토큰 갱신',
            'GET    /api/v1/auth/me            # 사용자 정보'
        ]),
        ('4.2 전략 API', [
            'GET    /api/v1/strategies         # 전략 목록',
            'POST   /api/v1/strategies         # 전략 생성',
            'GET    /api/v1/strategies/{id}    # 전략 상세',
            'PUT    /api/v1/strategies/{id}    # 전략 수정',
            'DELETE /api/v1/strategies/{id}    # 전략 삭제'
        ]),
        ('4.3 백테스팅 API', [
            'POST   /api/v1/backtests          # 백테스팅 시작',
            'GET    /api/v1/backtests          # 백테스팅 목록',
            'GET    /api/v1/backtests/{id}     # 백테스팅 결과',
            'GET    /api/v1/backtests/{id}/status  # 진행 상태',
            'DELETE /api/v1/backtests/{id}     # 백테스팅 삭제'
        ]),
        ('4.4 데이터 API', [
            'GET    /api/v1/data/symbols       # 심볼 목록',
            'GET    /api/v1/data/historical    # 과거 데이터',
            'POST   /api/v1/data/import        # 데이터 가져오기'
        ]),
        ('4.5 분석 API', [
            'GET    /api/v1/analytics/metrics  # 성과 지표',
            'GET    /api/v1/analytics/chart    # 차트 데이터',
            'GET    /api/v1/analytics/compare  # 전략 비교'
        ])
    ]

    for section_title, endpoints in api_sections:
        doc.add_heading(section_title, 2)
        p = doc.add_paragraph('\n'.join(endpoints))
        p.style = 'No Spacing'
        for run in p.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(9)

    # 5. 데이터베이스 스키마
    doc.add_heading('5. 데이터베이스 스키마', 1)

    db_tables = [
        ('5.1 Users 테이블', [
            'id (UUID, PK)',
            'email (VARCHAR, UNIQUE)',
            'username (VARCHAR, UNIQUE)',
            'password_hash (VARCHAR)',
            'created_at (TIMESTAMP)',
            'updated_at (TIMESTAMP)'
        ]),
        ('5.2 Strategies 테이블', [
            'id (UUID, PK)',
            'user_id (UUID, FK)',
            'name (VARCHAR)',
            'description (TEXT)',
            'code (TEXT) - 전략 코드 (Python)',
            'parameters (JSONB) - 파라미터',
            'is_public (BOOLEAN)',
            'created_at (TIMESTAMP)',
            'updated_at (TIMESTAMP)'
        ]),
        ('5.3 Backtests 테이블', [
            'id (UUID, PK)',
            'user_id (UUID, FK)',
            'strategy_id (UUID, FK)',
            'symbol (VARCHAR) - 종목',
            'start_date (DATE)',
            'end_date (DATE)',
            'initial_capital (DECIMAL)',
            'status (ENUM: pending, running, completed, failed)',
            'result (JSONB) - 백테스팅 결과',
            'created_at (TIMESTAMP)',
            'completed_at (TIMESTAMP)'
        ]),
        ('5.4 MarketData 테이블', [
            'id (BIGSERIAL, PK)',
            'symbol (VARCHAR)',
            'timestamp (TIMESTAMP)',
            'open (DECIMAL)',
            'high (DECIMAL)',
            'low (DECIMAL)',
            'close (DECIMAL)',
            'volume (DECIMAL)',
            'INDEX: (symbol, timestamp)'
        ]),
        ('5.5 BacktestTrades 테이블', [
            'id (BIGSERIAL, PK)',
            'backtest_id (UUID, FK)',
            'timestamp (TIMESTAMP)',
            'type (ENUM: buy, sell)',
            'price (DECIMAL)',
            'quantity (DECIMAL)',
            'commission (DECIMAL)'
        ])
    ]

    for table_title, fields in db_tables:
        doc.add_heading(table_title, 2)
        for field in fields:
            doc.add_paragraph(field, style='List Bullet')

    # 6. Fly.io 배포 전략
    doc.add_heading('6. Fly.io 배포 전략', 1)

    doc.add_heading('6.1 배포 아키텍처', 2)
    doc.add_paragraph('옵션 A: 단일 앱 (권장 - 비용 절감)', style='List Bullet')
    doc.add_paragraph('  - Nginx를 프록시로 사용', style='List Bullet')
    doc.add_paragraph('  - /api → Backend (FastAPI)', style='List Bullet')
    doc.add_paragraph('  - / → Frontend (Next.js)', style='List Bullet')
    doc.add_paragraph()
    doc.add_paragraph('옵션 B: 분리 앱 (확장성 우선)', style='List Bullet')
    doc.add_paragraph('  - Frontend와 Backend를 별도 Fly.io 앱으로 배포', style='List Bullet')
    doc.add_paragraph('  - 독립적인 스케일링 가능', style='List Bullet')

    doc.add_heading('6.2 환경 변수 관리', 2)
    env_mgmt = [
        'fly secrets set을 통한 시크릿 관리',
        '환경별 .env 파일 분리'
    ]
    for item in env_mgmt:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('6.3 볼륨 설정', 2)
    volume_items = [
        'PostgreSQL 데이터: Persistent Volume',
        'Redis 데이터: 메모리 (재시작시 초기화 OK)'
    ]
    for item in volume_items:
        doc.add_paragraph(item, style='List Bullet')

    # 7. 주요 화면 구성
    doc.add_heading('7. 주요 화면 구성', 1)

    screens = [
        ('7.1 홈페이지 (/)', [
            '서비스 소개',
            '주요 기능 소개',
            '로그인/회원가입 링크'
        ]),
        ('7.2 대시보드 (/dashboard)', [
            '최근 백테스팅 결과 요약',
            '전략 성과 개요',
            '빠른 실행 버튼'
        ]),
        ('7.3 전략 관리 (/strategies)', [
            '전략 목록 (테이블)',
            '전략 생성/수정 폼',
            '코드 에디터 (Monaco Editor)',
            '전략 테스트 기능'
        ]),
        ('7.4 백테스팅 실행 (/backtests/new)', [
            '전략 선택',
            '기간 설정 (시작일, 종료일)',
            '초기 자본 설정',
            '종목 선택',
            '실행 버튼'
        ]),
        ('7.5 백테스팅 결과 (/backtests/{id})', [
            '수익 곡선 차트',
            '주요 지표 (수익률, MDD, 샤프비율 등)',
            '거래 내역 테이블',
            '드로우다운 차트',
            '월별 수익률 히트맵'
        ]),
        ('7.6 분석 (/analytics)', [
            '여러 전략 비교',
            '성과 분석 대시보드',
            '커스텀 차트'
        ])
    ]

    for screen_title, features in screens:
        doc.add_heading(screen_title, 2)
        for feature in features:
            doc.add_paragraph(feature, style='List Bullet')

    # 8. 개발 워크플로우
    doc.add_heading('8. 개발 워크플로우', 1)

    doc.add_heading('8.1 로컬 개발 환경 설정', 2)
    local_setup = [
        '저장소 클론',
        'Docker Compose로 PostgreSQL, Redis 실행',
        'Backend: poetry install → uvicorn 실행',
        'Frontend: pnpm install → pnpm dev',
        'http://localhost:3000 접속'
    ]
    for i, step in enumerate(local_setup, 1):
        doc.add_paragraph(f'{i}. {step}')

    doc.add_heading('8.2 데이터베이스 마이그레이션', 2)
    migration_steps = [
        'alembic revision --autogenerate -m "message"',
        'alembic upgrade head'
    ]
    for i, step in enumerate(migration_steps, 1):
        doc.add_paragraph(f'{i}. {step}')

    doc.add_heading('8.3 배포 프로세스', 2)
    deploy_steps = [
        'Git push to main',
        'GitHub Actions CI 실행 (테스트)',
        'Docker 이미지 빌드',
        'Fly.io 배포 (flyctl deploy)',
        '마이그레이션 실행'
    ]
    for i, step in enumerate(deploy_steps, 1):
        doc.add_paragraph(f'{i}. {step}')

    # 9. 성능 최적화
    doc.add_heading('9. 성능 최적화', 1)

    doc.add_heading('9.1 백엔드', 2)
    backend_opt = [
        '데이터베이스 인덱싱',
        'Redis 캐싱 (자주 조회되는 데이터)',
        'Celery 비동기 작업 (백테스팅)',
        'Connection pooling (asyncpg)'
    ]
    for item in backend_opt:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('9.2 프론트엔드', 2)
    frontend_opt = [
        'Next.js SSR/SSG 활용',
        '이미지 최적화 (next/image)',
        '코드 스플리팅',
        'React Query 캐싱',
        '차트 데이터 페이지네이션'
    ]
    for item in frontend_opt:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('9.3 인프라', 2)
    infra_opt = [
        'Nginx gzip 압축',
        'CDN (정적 파일)',
        '데이터베이스 쿼리 최적화'
    ]
    for item in infra_opt:
        doc.add_paragraph(item, style='List Bullet')

    # 10. 보안 고려사항
    doc.add_heading('10. 보안 고려사항', 1)

    doc.add_heading('10.1 인증 & 인가', 2)
    auth_security = [
        'JWT 토큰 (Access + Refresh)',
        'HTTPS 강제',
        'CORS 설정',
        'Rate limiting'
    ]
    for item in auth_security:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('10.2 데이터 보호', 2)
    data_security = [
        '비밀번호 해싱 (bcrypt)',
        'SQL Injection 방지 (ORM 사용)',
        'XSS 방지 (입력 검증)',
        'CSRF 토큰'
    ]
    for item in data_security:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('10.3 API 보안', 2)
    api_security = [
        'API Key 관리',
        '요청 검증 (Pydantic)',
        '에러 메시지 최소화'
    ]
    for item in api_security:
        doc.add_paragraph(item, style='List Bullet')

    # 11. Supabase 마이그레이션 계획
    doc.add_heading('11. Supabase 마이그레이션 계획', 1)

    doc.add_heading('11.1 마이그레이션 단계', 2)
    migration_phases = [
        'Phase 1: Supabase 프로젝트 생성',
        'Phase 2: 스키마 복제 (Supabase SQL Editor)',
        'Phase 3: 데이터 마이그레이션 (pg_dump/restore)',
        'Phase 4: 연결 문자열 변경',
        'Phase 5: Supabase Auth 통합 (선택)'
    ]
    for phase in migration_phases:
        doc.add_paragraph(phase, style='List Bullet')

    doc.add_heading('11.2 Supabase 추가 기능 활용', 2)
    supabase_features = [
        'Supabase Auth: 소셜 로그인',
        'Supabase Storage: 파일 저장',
        'Realtime: 백테스팅 실시간 업데이트',
        'Edge Functions: 서버리스 함수'
    ]
    for feature in supabase_features:
        doc.add_paragraph(feature, style='List Bullet')

    # 12. 모니터링 & 로깅
    doc.add_heading('12. 모니터링 & 로깅', 1)

    doc.add_heading('12.1 로깅', 2)
    logging_items = [
        '구조화된 로깅 (JSON)',
        '로그 레벨 관리',
        'Fly.io logs 연동'
    ]
    for item in logging_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('12.2 에러 트래킹', 2)
    error_tracking = [
        'Sentry 통합 (백엔드/프론트엔드)',
        '에러 알림'
    ]
    for item in error_tracking:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('12.3 성능 모니터링', 2)
    perf_monitoring = [
        'API 응답 시간',
        '백테스팅 실행 시간',
        '데이터베이스 쿼리 성능'
    ]
    for item in perf_monitoring:
        doc.add_paragraph(item, style='List Bullet')

    # 13. 개발 일정
    doc.add_heading('13. 개발 일정 (추정)', 1)

    timeline = [
        ('Week 1-2: 환경 설정 및 기본 구조', [
            '모노레포 설정',
            'Docker 환경 구성',
            'DB 스키마 설계',
            '기본 API 구조'
        ]),
        ('Week 3-4: 백엔드 핵심 기능', [
            '인증 시스템',
            '백테스팅 엔진',
            'Celery 작업 큐',
            'API 엔드포인트'
        ]),
        ('Week 5-6: 프론트엔드 개발', [
            'UI 컴포넌트',
            '페이지 구현',
            'API 연동',
            '차트 시각화'
        ]),
        ('Week 7: 배포 및 테스팅', [
            'Fly.io 배포 설정',
            'CI/CD 파이프라인',
            '통합 테스트',
            '성능 최적화'
        ]),
        ('Week 8: 개선 및 문서화', [
            '버그 수정',
            '문서 작성',
            '사용자 가이드'
        ])
    ]

    for week_title, tasks in timeline:
        doc.add_heading(week_title, 2)
        for task in tasks:
            doc.add_paragraph(task, style='List Bullet')

    # 문서 정보
    doc.add_paragraph()
    doc.add_paragraph('=' * 60)
    footer = doc.add_paragraph('문서 작성일: 2025-11-15')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('=' * 60)

    # 저장
    doc.save('/home/jbwi127/trading-backtest/웹배포_기술스택.docx')
    print('✓ 웹배포 기술스택 DOCX 문서가 성공적으로 생성되었습니다: 웹배포_기술스택.docx')

if __name__ == '__main__':
    create_web_stack_doc()
